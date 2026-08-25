from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

# Diretórios
caminho_outros = Path(r'C:\TEMPORÁRIOS\OUTROS')
caminho_clientes = Path(r'C:\TEMPORÁRIOS\CLIENTES')

caminho_outros.mkdir(parents=True, exist_ok=True)
caminho_clientes.mkdir(parents=True, exist_ok=True)

# 1. Busca o relatório de pastas mais recente no diretório OUTROS
relatorios_pastas = [
    f
    for f in caminho_outros.glob('Relatorio_Pastas*.docx')
    if not f.name.startswith('~$')  # Ignora arquivos temporários do Word
]

if not relatorios_pastas:
    raise FileNotFoundError(
        "Nenhum arquivo 'Relatorio_Pastas*.docx' foi encontrado em C:\\TEMPORÁRIOS\\OUTROS."
    )

relatorio_mais_recente = max(
    relatorios_pastas, key=lambda f: f.stat().st_mtime
)

# 2. Extrai do Word os nomes de pastas marcadas como [CRIADA]
doc_base = Document(relatorio_mais_recente)
pastas_para_validar = []

for paragraph in doc_base.paragraphs:
    texto = paragraph.text.strip()
    if texto.startswith('[CRIADA]'):
        nome_pasta = texto.replace('[CRIADA]', '').strip()
        if nome_pasta:
            pastas_para_validar.append(nome_pasta)

# 3. Valida e cria a pasta "SALDO INICIAL" dentro de cada cliente
saldos_criados = []
saldos_ignorados = []

for nome_pasta in pastas_para_validar:
    pasta_cliente = caminho_clientes / nome_pasta
    pasta_saldo_inicial = pasta_cliente / 'SALDO INICIAL'

    if pasta_cliente.exists() and pasta_cliente.is_dir():
        if not pasta_saldo_inicial.exists():
            pasta_saldo_inicial.mkdir(parents=True, exist_ok=True)
            saldos_criados.append(nome_pasta)
        else:
            saldos_ignorados.append(nome_pasta)
    else:
        # Se a pasta principal não existir fisicamente no disco
        saldos_ignorados.append(
            f'{nome_pasta} (Pasta principal não encontrada)'
        )

# 4. Monta o relatório Word de Saldos Iniciais
doc = Document()

# Cabeçalho
doc.add_heading('Relatório de Validação de Saldos Iniciais', level=0)
doc.add_paragraph(
    f'Data de execução: {datetime.now().strftime("%d/%m/%Y às %H:%M:%S")}'
)
doc.add_paragraph(f'Relatório de origem: {relatorio_mais_recente.name}')

# Resumo
doc.add_heading('Resumo Geral', level=1)
p_resumo = doc.add_paragraph()
p_resumo.add_run('• Subpastas "SALDO INICIAL" criadas: ').bold = True
p_resumo.add_run(f'{len(saldos_criados)}\n')
p_resumo.add_run(
    '• Subpastas "SALDO INICIAL" ignoradas (já existentes): '
).bold = True
p_resumo.add_run(f'{len(saldos_ignorados)}')

# Seção de Saldos Criados
doc.add_heading(f'Subpastas Criadas ({len(saldos_criados)})', level=1)
if saldos_criados:
    for item in saldos_criados:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[CRIADA] {item}\\SALDO INICIAL')
        run.font.color.rgb = RGBColor(34, 139, 34)  # Verde
else:
    doc.add_paragraph('Nenhuma subpasta "SALDO INICIAL" precisou ser criada.')

# Seção de Saldos Ignorados
doc.add_heading(f'Subpastas Ignoradas ({len(saldos_ignorados)})', level=1)
if saldos_ignorados:
    for item in saldos_ignorados:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[IGNORADA] {item}\\SALDO INICIAL')
        run.font.color.rgb = RGBColor(204, 102, 0)  # Laranja
else:
    doc.add_paragraph('Nenhuma subpasta foi ignorada.')

# Salva o arquivo no diretório OUTROS
nome_relatorio = (
    f'Relatorio_Saldos_Iniciais_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
)
caminho_final = caminho_outros / nome_relatorio
doc.save(caminho_final)

print('Processo concluído!')
print(f'Relatório salvo em: {caminho_final}')