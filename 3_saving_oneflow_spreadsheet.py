import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

# Diretórios
caminho_clientes = Path(r'C:\TEMPORÁRIOS\CLIENTES')
caminho_outros = Path(r'C:\TEMPORÁRIOS\OUTROS')

caminho_outros.mkdir(parents=True, exist_ok=True)

# Arquivo de origem (Modelo)
arquivo_modelo = (
    caminho_clientes
    / '0000 - MODELO'
    / 'SALDO INICIAL'
    / 'oneflowsaldoinicial.xlsx'
)

if not arquivo_modelo.exists():
    raise FileNotFoundError(
        f'O arquivo modelo não foi encontrado em: {arquivo_modelo}'
    )

copiados = []
ignorados = []

# 1. Varre todas as pastas de empresas (ignorando a pasta de modelo)
for pasta_empresa in caminho_clientes.iterdir():
    if pasta_empresa.is_dir() and pasta_empresa.name != '0000 - MODELO':
        pasta_saldo = pasta_empresa / 'SALDO INICIAL'

        # Verifica se a empresa possui a pasta SALDO INICIAL
        if pasta_saldo.exists() and pasta_saldo.is_dir():
            nome_empresa = pasta_empresa.name
            nome_novo_arquivo = f'oneflowsaldoinicial - {nome_empresa}.xlsx'
            caminho_destino = pasta_saldo / nome_novo_arquivo

            # Copia apenas se o arquivo ainda não existir no destino
            if not caminho_destino.exists():
                shutil.copy2(arquivo_modelo, caminho_destino)
                copiados.append({
                    'empresa': nome_empresa,
                    'arquivo': nome_novo_arquivo,
                })
            else:
                ignorados.append({
                    'empresa': nome_empresa,
                    'arquivo': nome_novo_arquivo,
                })

# 2. Monta o relatório Word
doc = Document()

# Cabeçalho
doc.add_heading('Relatório de Cópia - OneFlow Saldo Inicial', level=0)
doc.add_paragraph(
    f'Data de execução: {datetime.now().strftime("%d/%m/%Y às %H:%M:%S")}'
)

# Resumo
doc.add_heading('Resumo Geral', level=1)
p_resumo = doc.add_paragraph()
p_resumo.add_run('• Arquivos copiados com sucesso: ').bold = True
p_resumo.add_run(f'{len(copiados)}\n')
p_resumo.add_run('• Arquivos ignorados (já existentes): ').bold = True
p_resumo.add_run(f'{len(ignorados)}')

# Seção de Copiados
doc.add_heading(f'Arquivos Copiados ({len(copiados)})', level=1)
if copiados:
    for item in copiados:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[COPIADO] {item["empresa"]}\\{item["arquivo"]}')
        run.font.color.rgb = RGBColor(34, 139, 34)  # Verde
else:
    doc.add_paragraph('Nenhum arquivo novo precisou ser copiado.')

# Seção de Ignorados
doc.add_heading(f'Arquivos Ignorados ({len(ignorados)})', level=1)
if ignorados:
    for item in ignorados:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[IGNORADO] {item["empresa"]}\\{item["arquivo"]}')
        run.font.color.rgb = RGBColor(204, 102, 0)  # Laranja
else:
    doc.add_paragraph('Nenhum arquivo foi ignorado nesta execução.')

# Salva o relatório no diretório OUTROS
nome_relatorio = (
    f'Relatorio_OneFlow_Inicio_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
)
caminho_final = caminho_outros / nome_relatorio
doc.save(caminho_final)

print('Processo concluído com sucesso!')
print(f'Relatório salvo em: {caminho_final}')