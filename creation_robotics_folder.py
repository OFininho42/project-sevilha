from datetime import datetime
from pathlib import Path
import re
import unicodedata
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl as opxl


def limpar_nome_pasta(nome):
    """Remove caracteres proibidos no Windows para nomear pastas."""
    return re.sub(r'[\\/*?:"<>|]', '', str(nome)).strip()


def normalizar_para_comparacao(texto):
    """Remove acentos, símbolos e espaços para comparar equivalência de nomes."""
    texto_sem_acento = (
        unicodedata.normalize('NFKD', str(texto))
        .encode('ASCII', 'ignore')
        .decode('utf-8')
    )
    return re.sub(r'[^a-zA-Z0-9]', '', texto_sem_acento).lower()


class Controle:
    wb_path = r'C:\TEMPORÁRIOS\IMPORTANTE\Meu_Controle.xlsx'
    wb = opxl.load_workbook(wb_path)
    ws = wb['Planilha1']

    codempresa = 'A'
    nomeempresa = 'B'


meu_controle = Controle()

# Diretórios
caminho_clientes = Path(r'C:\TEMPORÁRIOS\CLIENTES')
caminho_outros = Path(r'C:\TEMPORÁRIOS\OUTROS')

caminho_clientes.mkdir(parents=True, exist_ok=True)
caminho_outros.mkdir(parents=True, exist_ok=True)

# 1. Mapeia pastas existentes no disco
pastas_existentes_norm = {
    normalizar_para_comparacao(p.name): p.name
    for p in caminho_clientes.iterdir()
    if p.is_dir()
}

# Listas para consolidar o relatório
criadas = []
ignoradas = []

# 2. Processa a planilha
for linha in range(2, meu_controle.ws.max_row + 1):
    val_a = meu_controle.ws[f'{meu_controle.codempresa}{linha}'].value
    val_b = meu_controle.ws[f'{meu_controle.nomeempresa}{linha}'].value

    if val_a is not None or val_b is not None:
        val_a_str = str(val_a).strip() if val_a is not None else ''
        val_b_str = str(val_b).strip() if val_b is not None else ''

        nome_bruto = f'{val_a_str} - {val_b_str}'
        nome_limpo = limpar_nome_pasta(nome_bruto)
        chave_excel = normalizar_para_comparacao(nome_bruto)

        if chave_excel:
            if chave_excel not in pastas_existentes_norm:
                # Criar pasta
                caminho_nova_pasta = caminho_clientes / nome_limpo
                caminho_nova_pasta.mkdir(parents=True, exist_ok=True)

                pastas_existentes_norm[chave_excel] = nome_limpo
                criadas.append(nome_limpo)
            else:
                pasta_atual = pastas_existentes_norm.get(
                    chave_excel, 'Já existente'
                )
                ignoradas.append({
                    'tentativa': nome_limpo,
                    'existente': pasta_atual,
                })

# 3. Monta o arquivo Word (.docx)
doc = Document()

# Título e Cabeçalho
doc.add_heading('Relatório de Processamento de Pastas', level=0)
doc.add_paragraph(
    f'Data de execução: {datetime.now().strftime("%d/%m/%Y às %H:%M:%S")}'
)

# Resumo
doc.add_heading('Resumo Geral', level=1)
p_resumo = doc.add_paragraph()
p_resumo.add_run('• Total de pastas criadas: ').bold = True
p_resumo.add_run(f'{len(criadas)}\n')
p_resumo.add_run('• Total de pastas ignoradas (já existentes): ').bold = True
p_resumo.add_run(f'{len(ignoradas)}')

# Seção de Pastas Criadas
doc.add_heading(f'Pastas Criadas ({len(criadas)})', level=1)
if criadas:
    for item in criadas:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[CRIADA] {item}')
        run.font.color.rgb = RGBColor(34, 139, 34)  # Verde
else:
    doc.add_paragraph('Nenhuma pasta nova foi criada nesta execução.')

# Seção de Pastas Ignoradas
doc.add_heading(f'Pastas Ignoradas ({len(ignoradas)})', level=1)
if ignoradas:
    for item in ignoradas:
        p = doc.add_paragraph(style='List Bullet')
        run_status = p.add_run(f'[IGNORADA] {item["tentativa"]}')
        run_status.font.color.rgb = RGBColor(204, 102, 0)  # Laranja
        p.add_run(f'  →  (Correspondente no disco: "{item["existente"]}")')
else:
    doc.add_paragraph('Nenhuma pasta foi ignorada nesta execução.')

# Salva em C:\TEMPORÁRIOS\OUTROS com timestamp
nome_relatorio = (
    f'Relatorio_Pastas_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
)
caminho_final_doc = caminho_outros / nome_relatorio
doc.save(caminho_final_doc)

print('Processo concluído com sucesso!')
print(f'Relatório salvo em: {caminho_final_doc}')